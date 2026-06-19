#!/usr/bin/env python3
"""Hybrid grader for written-report assignments (A4-style).

Architecture:
  Python (this file): extract text from PDF/DOCX/MD, compute structural signals
                      (sections, length, APA citations, image count).
  Agent (LLM):        reads the extracted text, applies rubric criteria that
                      need judgment (depth, originality of conclusions,
                      relevance of examples, APA correctness), and writes the
                      combined findings via lib/grading_md.

Usage:
    python3 grade_informe_template.py <folder> [--slug asignacion-04-informe-grupal]

The folder structure expected:
    <folder>/<student_or_group_name>/<one or more .docx|.pdf|.md files>

The grader processes each subfolder as one submission. For group work where
N students share one submission, the folder layout can be:
    <folder>/GROUP_<topic>/<file>.docx
    <folder>/GROUP_<topic>/MEMBERS.txt   # one student folder name per line

The grader emits per-group `STRUCTURAL.json` with extracted text + signals.
The agent then reads each STRUCTURAL.json, judges the qualitative criteria,
and calls write_grading_md with the combined `findings` dict.
"""
from __future__ import annotations
import json
import re
import sys
from pathlib import Path


# ─────────────────────────────────────────────────────────────────────────────
# Extraction
# ─────────────────────────────────────────────────────────────────────────────

def extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise SystemExit("python-docx not installed. `pip install python-docx`")
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # Tables
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def count_docx_images(path: Path) -> int:
    """Count embedded images in a .docx — they live in word/media/ inside the zip."""
    import zipfile
    try:
        with zipfile.ZipFile(str(path)) as z:
            return sum(1 for n in z.namelist() if n.startswith("word/media/"))
    except Exception:
        return 0


def extract_pdf(path: Path) -> str:
    # Try pdfplumber first, fall back to pdftotext CLI
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)
    except ImportError:
        pass
    import subprocess
    r = subprocess.run(["pdftotext", "-layout", str(path), "-"], capture_output=True, text=True)
    if r.returncode != 0:
        raise RuntimeError(f"pdftotext failed: {r.stderr[:200]}")
    return r.stdout


def count_pdf_images(path: Path) -> int:
    """Heuristic: count `/Image` XObject occurrences via pdftotext is impossible.
    Use pdfimages -list if available."""
    import subprocess
    r = subprocess.run(["pdfimages", "-list", str(path)], capture_output=True, text=True)
    if r.returncode != 0:
        return 0
    # Skip header lines, count data rows
    return max(0, len(r.stdout.strip().splitlines()) - 2)


def extract_md(path: Path) -> str:
    return path.read_text(errors="ignore")


def extract_text(file: Path) -> tuple[str, int]:
    """Returns (text, image_count). Detects format by suffix."""
    suf = file.suffix.lower()
    if suf == ".docx":
        return extract_docx(file), count_docx_images(file)
    if suf == ".pdf":
        return extract_pdf(file), count_pdf_images(file)
    if suf in (".md", ".txt"):
        return extract_md(file), 0
    raise ValueError(f"Unsupported format: {suf}")


# ─────────────────────────────────────────────────────────────────────────────
# Structural signals (deterministic)
# ─────────────────────────────────────────────────────────────────────────────

SECTION_KEYWORDS = {
    "portada": [r"^portada$", r"^cover$", r"título.*tema", r"^introducci[oó]n", r"\bintroducci[oó]n\b"],
    "contenido": [r"\bdesarrollo\b", r"\bcontenido\b", r"\banálisis\b", r"\banalisis\b"],
    "conclusiones": [r"\bconclusiones?\b", r"\bconclusion\b"],
    "bibliografia": [r"\bbibliograf[ií]a\b", r"\breferencias?\b", r"\bfuentes?\b"],
}


def analyze_text(text: str, image_count: int) -> dict:
    lines = text.splitlines()
    nonblank = [l for l in lines if l.strip()]
    words = re.findall(r"\b\w+\b", text)

    # Sections present (case-insensitive)
    sections = {}
    for key, pats in SECTION_KEYWORDS.items():
        found = False
        for line in lines:
            for pat in pats:
                if re.search(pat, line, re.IGNORECASE):
                    found = True
                    break
            if found:
                break
        sections[key] = found

    # APA citations: (Author, Year) or (Author et al., Year)
    apa_inline = re.findall(r"\(([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+(?:et\s+al\.|and|y)\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)?,\s*\d{4}[a-z]?)\)", text)
    apa_inline_count = len(apa_inline)

    # References: look for the bibliography section and count entries (lines with Year)
    bib_start = None
    for i, line in enumerate(lines):
        if re.search(r"\bbibliograf[ií]a|referencias?\b", line, re.IGNORECASE):
            bib_start = i
            break
    refs_count = 0
    if bib_start is not None:
        # Crude: count lines after bib_start that look like an APA ref (contains a year)
        for line in lines[bib_start + 1:]:
            if re.search(r"\(\d{4}\)|,\s*\d{4}\.", line):
                refs_count += 1

    # Conclusion personal markers: detect named conclusions like "Conclusión de Juan"
    personal_concl_markers = len(re.findall(r"conclusi[oó]n\s+(?:de|por)\s+[A-ZÁÉÍÓÚÑ][\wáéíóúñ]+", text, re.IGNORECASE))

    return {
        "char_count": len(text),
        "word_count": len(words),
        "line_count": len(nonblank),
        "image_count": image_count,
        "sections": sections,
        "sections_present_count": sum(1 for v in sections.values() if v),
        "apa_inline_citations": apa_inline_count,
        "apa_references_in_bib": refs_count,
        "personal_conclusions_markers": personal_concl_markers,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Structural scoring (partial — only the criteria that don't need LLM)
# ─────────────────────────────────────────────────────────────────────────────

def score_structural(signals: dict) -> dict:
    """Returns levels 1-4 for the criteria that can be scored mechanically.
    The agent (LLM) fills in `profundidad` and `conclusiones_personales` after
    reading the extracted text."""
    # Estructura organizada — based on sections present
    sec_n = signals["sections_present_count"]
    estructura = {4: 4, 3: 3, 2: 2, 1: 1, 0: 1}.get(sec_n, 1)

    # Formato APA — based on inline citations + bib refs
    refs = signals["apa_references_in_bib"]
    inline = signals["apa_inline_citations"]
    if refs >= 5 and inline >= 5:
        apa = 4
    elif refs >= 3 and inline >= 3:
        apa = 3
    elif refs >= 1 or inline >= 1:
        apa = 2
    else:
        apa = 1

    # Ejemplos visuales — based on image count (LLM judges relevance later)
    images = signals["image_count"]
    if images >= 5:
        ejemplos_visuales_structural = 4
    elif images >= 3:
        ejemplos_visuales_structural = 3
    elif images >= 1:
        ejemplos_visuales_structural = 2
    else:
        ejemplos_visuales_structural = 1

    return {
        "estructura": estructura,
        "formato_apa_structural": apa,
        "ejemplos_visuales_structural": ejemplos_visuales_structural,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Rubric metadata (shared with build_findings)
# ─────────────────────────────────────────────────────────────────────────────

CRITERIA = [
    # key, label, weight_pct, judged_by
    ("profundidad", "Profundidad del contenido", 20, "llm"),
    ("conclusiones_personales", "Conclusiones personales", 20, "llm"),
    ("formato_apa", "Formato APA correcto", 20, "hybrid"),  # structural floor + LLM verify
    ("ejemplos_visuales", "Ejemplos visuales", 20, "hybrid"),  # count + LLM relevance
    ("estructura", "Estructura organizada", 20, "structural"),
]


def build_findings(signals: dict, llm_scores: dict, recovered_from: str | None = None) -> dict:
    """Combine structural signals + LLM-judged scores into the standard findings dict.

    `llm_scores` must include: profundidad, conclusiones_personales, formato_apa, ejemplos_visuales
    (the agent fills these after reading the extracted text).
    `signals["estructura"]` from score_structural() is used for the structural criterion.
    """
    structural = score_structural(signals) if "estructura" not in signals else signals
    levels = {
        "profundidad": llm_scores["profundidad"],
        "conclusiones_personales": llm_scores["conclusiones_personales"],
        "formato_apa": llm_scores.get("formato_apa", structural["formato_apa_structural"]),
        "ejemplos_visuales": llm_scores.get("ejemplos_visuales", structural["ejemplos_visuales_structural"]),
        "estructura": structural["estructura"],
    }
    rubric = []
    total_pts = 0
    for key, label, weight, _ in CRITERIA:
        lvl = levels[key]
        rubric.append({
            "criterion": label,
            "level": lvl,
            "max": 4,
            "weight_pct": weight,
            "note": {4: "Excelente", 3: "Bueno", 2: "Normal", 1: "Deficiente"}.get(lvl, "?"),
        })
        # weight 20% → each level pt worth (20/4)=5 of the 100 total
        total_pts += lvl * (weight / 4)
    score = round(total_pts)
    missing = [r["criterion"] for r in rubric if r["level"] <= 2]
    return {
        "score": score,
        "rubric": rubric,
        "signals": signals,
        "missing": missing,
        "recovered_from": recovered_from,
    }


# ─────────────────────────────────────────────────────────────────────────────
# CLI: extract + structural signals for one folder (per submission)
# ─────────────────────────────────────────────────────────────────────────────

def process_folder(folder: Path, slug: str, out_dir: Path) -> dict | None:
    """Process one submission folder. Returns metadata dict or None if no file found."""
    files = [f for f in folder.iterdir()
             if f.is_file() and f.suffix.lower() in (".docx", ".pdf", ".md", ".txt")]
    if not files:
        return None
    file = files[0]  # use first if multiple
    try:
        text, images = extract_text(file)
    except Exception as e:
        return {"folder": folder.name, "error": str(e)}
    signals = analyze_text(text, images)
    structural = score_structural(signals)
    # Persist extracted text + signals so the agent can read them in Phase 4
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / f"{folder.name}.txt").write_text(text)
    payload = {
        "folder": folder.name,
        "file": file.name,
        "signals": signals,
        "structural_levels": structural,
        "ready_for_llm": True,
    }
    (out_dir / f"{folder.name}.json").write_text(json.dumps(payload, indent=2, ensure_ascii=False))
    return payload


def main():
    if len(sys.argv) < 2:
        print("Usage: grade_informe_template.py <folder-with-submissions> [--slug asignacion-04-informe-grupal]")
        sys.exit(1)
    target = Path(sys.argv[1])
    slug = "asignacion-04-informe-grupal"
    if "--slug" in sys.argv:
        slug = sys.argv[sys.argv.index("--slug") + 1]

    submissions_root = Path("/Users/ea/Projects/universidad/regular/2026/SEMI/submissions") / slug
    out_dir = submissions_root / "extracted"

    # If target IS the submissions folder, iterate subfolders. Otherwise treat target as one.
    if target.is_file():
        print("Pass a folder, not a single file.", file=sys.stderr)
        sys.exit(1)
    children = [d for d in target.iterdir() if d.is_dir()]
    if not children:
        children = [target]

    rows = []
    for d in sorted(children):
        result = process_folder(d, slug, out_dir)
        if result:
            rows.append(result)
            sig = result.get("signals", {})
            print(f"{d.name:<35} words={sig.get('word_count', 0):<6} "
                  f"imgs={sig.get('image_count', 0):<3} secs={sig.get('sections_present_count', 0)}/4  "
                  f"apa_inline={sig.get('apa_inline_citations', 0)} refs={sig.get('apa_references_in_bib', 0)}")

    print(f"\nExtracted {len(rows)} submissions → {out_dir}")
    print("Next step: agent reads each .txt + .json in extracted/, judges profundidad/conclusiones/apa/ejemplos,")
    print("then calls grading_md.write_grading_md with build_findings(signals, llm_scores).")


if __name__ == "__main__":
    main()
